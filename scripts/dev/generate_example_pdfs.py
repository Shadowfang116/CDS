#!/usr/bin/env python3
"""
Generate safe placeholder PDFs and DOCX for docs/pilot_samples_real_example/
These files contain no PII and are safe to commit.
"""
import sys
from pathlib import Path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from io import BytesIO

# Add /app to path (when running in container)
sys.path.insert(0, '/app')

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def create_placeholder_pdf(title: str, content_lines: list, output_path: str):
    """Create a simple PDF with text."""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, title)
    
    # Content
    c.setFont("Helvetica", 12)
    y = height - 100
    for line in content_lines:
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line[:80])
        y -= 20
    
    c.save()
    buffer.seek(0)
    pdf_bytes = buffer.read()
    
    with open(output_path, 'wb') as f:
        f.write(pdf_bytes)
    
    print(f"Created: {output_path}")

def create_placeholder_docx(title: str, content_lines: list, output_path: str):
    """Create a simple DOCX with text."""
    if not HAS_DOCX:
        print(f"[WARN] python-docx not available, skipping DOCX: {output_path}")
        return
    
    doc = DocxDocument()
    
    # Title
    title_para = doc.add_heading(title, level=1)
    
    # Content
    for line in content_lines:
        para = doc.add_paragraph(line)
        para.paragraph_format.space_after = Pt(6)
    
    doc.save(output_path)
    print(f"Created: {output_path}")


def main():
    output_dir = Path("docs/pilot_samples_real_example")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Example DHA NDC
    create_placeholder_pdf(
        "Example DHA NDC Document",
        [
            "This is a placeholder DHA NDC document.",
            "Property Details:",
            "Plot Number: EXAMPLE-123",
            "Block: A",
            "Phase: 5",
            "Scheme: DHA",
            "This document is for demonstration purposes only.",
            "No real property information is included.",
        ],
        str(output_dir / "EXAMPLE_CASE__DHA_NDC__PLACEHOLDER.pdf")
    )
    
    # Example Sale Deed
    create_placeholder_pdf(
        "Example Sale Deed Document",
        [
            "This is a placeholder Sale Deed document.",
            "Registry Information:",
            "Registry Number: EXAMPLE-456/2023",
            "Registry Date: 2023-01-15",
            "Registry Office: LDA Lahore",
            "Instrument: Sale Deed",
            "This document is for demonstration purposes only.",
            "No real property information is included.",
        ],
        str(output_dir / "EXAMPLE_CASE__SALE_DEED__PLACEHOLDER.pdf")
    )
    
    # Example Fard
    create_placeholder_pdf(
        "Example Fard Document",
        [
            "This is a placeholder Fard document.",
            "Revenue Information:",
            "Khasra Number: EXAMPLE-789",
            "Khewat Number: EXAMPLE-012",
            "Mouza: Example Village",
            "Tehsil: Example Tehsil",
            "District: Example District",
            "This document is for demonstration purposes only.",
            "No real property information is included.",
        ],
        str(output_dir / "EXAMPLE_CASE__FARD__PLACEHOLDER.pdf")
    )
    
    # P13: Example Legal Opinion DOCX
    create_placeholder_docx(
        "PILOT DEMO OPINION",
        [
            "Legal Opinion for Property Transfer",
            "",
            "Property Details:",
            "Plot Number: DEMO-001",
            "Block: B",
            "Phase: 6",
            "Scheme: DHA",
            "",
            "Registry Information:",
            "Registry Number: DEMO-789/2024",
            "Registry Date: 2024-01-20",
            "Registry Office: LDA Lahore",
            "Instrument: Sale Deed",
            "",
            "Party Information:",
            "Borrower: Example Borrower Name",
            "Seller: Example Seller Name",
            "",
            "This is a placeholder legal opinion document for pilot testing.",
            "No real property information is included.",
            "This document is safe to commit to the repository.",
        ],
        str(output_dir / "PILOT_DEMO_OPINION.docx")
    )
    
    print("\n[OK] All placeholder files created successfully!")

if __name__ == "__main__":
    main()

