"""D5: DOCX draft generation service."""
import io
from datetime import datetime
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _add_header(doc: Document, org_name: str, case_ref: str, date: str, title: str):
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = header.add_run(org_name.upper())
    org_run.bold = True
    org_run.font.size = Pt(14)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(12)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("[DRAFT]")
    sub_run.bold = True
    sub_run.font.size = Pt(10)

    doc.add_paragraph(f"Matter Reference: {case_ref}")
    doc.add_paragraph(f"Date: {date}")
    doc.add_paragraph()


def _add_heading(doc: Document, title: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(11)


def _numbered_items(doc: Document, items: List[str], empty_text: str):
    if not items:
        doc.add_paragraph(f"1. {empty_text}")
        return

    for index, item in enumerate(items, start=1):
        doc.add_paragraph(f"{index}. {item}")


def _format_exception(exc: Dict[str, Any]) -> str:
    status = exc.get("status", "Open")
    text = f"[{exc.get('severity', 'Medium')}] {exc.get('title', 'Exception')}"
    if exc.get("description"):
        text += f": {exc['description']}"
    if status != "Open":
        text += f" ({status})"
    return text


def _format_cp(cp: Dict[str, Any]) -> str:
    status = cp.get("status", "Open")
    label = "Fulfilled" if status == "Satisfied" else ("Pending" if status == "Open" else status)
    text = f"[{cp.get('severity', 'Medium')}] {cp.get('text', '')}"
    if cp.get("evidence_required"):
        text += f" (Evidence: {cp['evidence_required']})"
    text += f" ({label})"
    return text


def _get_borrower_info(dossier: Dict[str, List[str]]) -> Dict[str, str]:
    info = {
        "name": "[DRAFT - Confirm Borrower Name]",
        "cnic": "[DRAFT - Confirm CNIC]",
        "address": "[DRAFT - Confirm Address]",
    }

    names = dossier.get("party.name.borrower", []) or dossier.get("party.name.raw", [])
    if names:
        info["name"] = names[0]

    cnics = dossier.get("party.cnic", [])
    if cnics:
        info["cnic"] = cnics[0]

    addresses = dossier.get("property.address", [])
    if addresses:
        info["address"] = addresses[0]

    return info


def _first_value(dossier: Dict[str, List[str]], *keys: str, default: str = "[DRAFT - Confirm]") -> str:
    for key in keys:
        values = dossier.get(key, [])
        if values:
            return values[0]
    return default


def _get_property_info(dossier: Dict[str, List[str]]) -> str:
    property_address = _first_value(dossier, "property.address", default="[DRAFT - Confirm Property Address]")
    scheme = _first_value(dossier, "property.scheme_name", default="")
    block = _first_value(dossier, "property.block", default="")
    parts = [property_address]
    if block and block not in property_address:
        parts.append(block)
    if scheme and scheme not in property_address:
        parts.append(scheme)
    return ", ".join(part for part in parts if part)


def generate_discrepancy_letter(
    case: Dict[str, Any],
    org: Dict[str, Any],
    dossier: Dict[str, List[str]],
    exceptions: List[Dict[str, Any]],
    cps: List[Dict[str, Any]],
) -> Tuple[bytes, str]:
    doc = Document()

    case_ref = str(case.get("id", ""))[:8].upper()
    date_str = datetime.utcnow().strftime("%d %B %Y")
    org_name = org.get("name", "Bank Diligence Platform")
    borrower = _get_borrower_info(dossier)
    property_desc = _get_property_info(dossier)

    _add_header(doc, org_name, case_ref, date_str, "DISCREPANCY LETTER")

    doc.add_paragraph(f"To: {borrower['name']}")
    doc.add_paragraph(f"Address: {borrower['address']}")
    doc.add_paragraph()

    subject = doc.add_paragraph()
    subject.add_run("Subject: ").bold = True
    subject.add_run(f"Outstanding documentary discrepancies — {property_desc}")
    doc.add_paragraph()

    doc.add_paragraph(
        "Dear Sir / Madam,\n\n"
        "We refer to the property-backed finance matter noted above and the documents presently made available for legal diligence. "
        "Upon review, the following discrepancies and outstanding documentary requirements require attention before the matter can be presented for approval."
    )
    doc.add_paragraph()

    _add_heading(doc, "Numbered Discrepancies")
    open_exceptions = [_format_exception(e) for e in exceptions if e.get("status") == "Open"]
    _numbered_items(doc, open_exceptions, "No open Exceptions are presently recorded in this draft.")
    doc.add_paragraph()

    _add_heading(doc, "Related Conditions Precedent")
    open_cps = [_format_cp(c) for c in cps if c.get("status") == "Open"]
    _numbered_items(doc, open_cps, "No outstanding Conditions Precedent are presently recorded in this draft.")
    doc.add_paragraph()

    doc.add_paragraph(
        "You are requested to provide the above clarifications, supporting evidence, and corrective documents at the earliest opportunity so that legal review may be progressed without avoidable delay. "
        "Nothing in this draft should be read as a legal conclusion, final clearance, or Waiver of any documentary requirement."
    )
    doc.add_paragraph()
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph()
    doc.add_paragraph("_____________________________")
    doc.add_paragraph("Authorized Signatory")
    doc.add_paragraph(org_name)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"DISCREPANCY_LETTER__CASE_{str(case.get('id', ''))}__{datetime.utcnow().strftime('%Y%m%d')}__v1.docx"
    return buffer.getvalue(), filename


def generate_undertaking_indemnity(
    case: Dict[str, Any],
    org: Dict[str, Any],
    dossier: Dict[str, List[str]],
    exceptions: List[Dict[str, Any]],
    cps: List[Dict[str, Any]],
) -> Tuple[bytes, str]:
    doc = Document()

    case_ref = str(case.get("id", ""))[:8].upper()
    date_str = datetime.utcnow().strftime("%d %B %Y")
    org_name = org.get("name", "Bank Diligence Platform")
    borrower = _get_borrower_info(dossier)
    property_desc = _get_property_info(dossier)

    _add_header(doc, org_name, case_ref, date_str, "UNDERTAKING / INDEMNITY")

    _add_heading(doc, "1. Parties")
    doc.add_paragraph(f"(a) {org_name} (the \"Bank / Instructing Institution\"); and")
    doc.add_paragraph(f"(b) {borrower['name']} bearing CNIC / registration reference {borrower['cnic']} (the \"Obligor\").")
    doc.add_paragraph()

    _add_heading(doc, "2. Recitals")
    doc.add_paragraph(f"(a) The Obligor has requested financing in relation to {property_desc}.")
    doc.add_paragraph("(b) Certain Exceptions and documentary matters remain subject to completion, clarification, or controlled Waiver.")
    doc.add_paragraph("(c) This instrument is provided as a draft structure only and remains subject to legal review, transaction specifics, and approver instructions. [DRAFT]")
    doc.add_paragraph()

    _add_heading(doc, "3. Operative Clauses")
    clause_items = [
        "[DRAFT] The Obligor undertakes to provide all remaining title, authority, and revenue record evidence required by the Bank within the period specified by the final transaction documents.",
        "[DRAFT] The Obligor confirms that all statements and documents delivered for the diligence exercise are true, complete, and not misleading in any material respect.",
        "[DRAFT] The Obligor agrees to indemnify the Bank and its legal advisers against loss arising from any inaccuracy, non-disclosure, or defect in title subsequently discovered, to the extent finally settled in the executed instrument.",
    ]
    waived_exceptions = [e for e in exceptions if e.get("status") == "Waived"]
    for waived_exc in waived_exceptions:
        clause_items.append(
            f"[DRAFT] Specific covenant relating to the Waiver of '{waived_exc.get('title', 'Exception')}' to be settled by transaction counsel before execution."
        )
    _numbered_items(doc, clause_items, "[DRAFT] Operative clauses to be inserted by transaction counsel.")
    doc.add_paragraph()

    _add_heading(doc, "4. Execution")
    doc.add_paragraph("Executed as a draft form for internal legal review only. [DRAFT]")
    doc.add_paragraph()
    doc.add_paragraph("_____________________________")
    doc.add_paragraph(f"Name: {borrower['name']}")
    doc.add_paragraph(f"CNIC / Registration: {borrower['cnic']}")
    doc.add_paragraph()
    doc.add_paragraph("Witness 1: _____________________")
    doc.add_paragraph("Witness 2: _____________________")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"UNDERTAKING_INDEMNITY__CASE_{str(case.get('id', ''))}__{datetime.utcnow().strftime('%Y%m%d')}__v1.docx"
    return buffer.getvalue(), filename


def generate_internal_opinion_skeleton(
    case: Dict[str, Any],
    org: Dict[str, Any],
    dossier: Dict[str, List[str]],
    exceptions: List[Dict[str, Any]],
    cps: List[Dict[str, Any]],
) -> Tuple[bytes, str]:
    doc = Document()

    case_ref = str(case.get("id", ""))[:8].upper()
    date_str = datetime.utcnow().strftime("%d %B %Y")
    org_name = org.get("name", "Bank Diligence Platform")
    borrower = _get_borrower_info(dossier)
    property_desc = _get_property_info(dossier)

    _add_header(doc, org_name, case_ref, date_str, "LEGAL OPINION SKELETON")

    open_critical_or_high = [
        e for e in exceptions if e.get("status") == "Open" and e.get("severity") in {"Critical", "High"}
    ]
    opinion_label = "Qualified / subject to noted Exceptions"
    if not open_critical_or_high and not [c for c in cps if c.get("status") == "Open"]:
        opinion_label = "[DRAFT] Subject to final confirmation of the complete matter record"

    _add_heading(doc, "1. Instruction Summary")
    doc.add_paragraph(
        f"We are instructed to review the legal diligence position for {borrower['name']} in connection with {property_desc}. "
        "This document is a structured draft only and is not intended to state final legal conclusions. [DRAFT]"
    )
    doc.add_paragraph()

    _add_heading(doc, "2. Scope of Review")
    doc.add_paragraph(
        "The scope of review should cover the title record, relevant authority permissions, revenue documentation, chargeability, and the documentary basis of any Exception or Condition Precedent recorded on the matter. "
        "Final scope language should be settled by supervising counsel for the relevant transaction. [DRAFT]"
    )
    doc.add_paragraph()

    _add_heading(doc, "3. Key Findings")
    finding_items = [_format_exception(exception) for exception in exceptions]
    _numbered_items(doc, finding_items, "[DRAFT] Key findings to be inserted following full legal review.")
    doc.add_paragraph()

    _add_heading(doc, "4. Opinion")
    doc.add_paragraph(
        f"Opinion status: {opinion_label}. This section should state the final legal view only after the reviewing lawyer has considered the complete title chain, authority record, and outstanding Conditions Precedent. [DRAFT]"
    )
    if cps:
        doc.add_paragraph("Outstanding Conditions Precedent relevant to the opinion include:")
        for cp in cps:
            doc.add_paragraph(f"- {_format_cp(cp)}")
    doc.add_paragraph()

    _add_heading(doc, "5. Signature Block")
    doc.add_paragraph("Prepared by: __________________________ [DRAFT]")
    doc.add_paragraph("Designation: __________________________")
    doc.add_paragraph(f"Date: {date_str}")
    doc.add_paragraph()
    doc.add_paragraph("Reviewed by: __________________________ [DRAFT]")
    doc.add_paragraph("Designation: __________________________")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"INTERNAL_OPINION__CASE_{str(case.get('id', ''))}__{datetime.utcnow().strftime('%Y%m%d')}__v1.docx"
    return buffer.getvalue(), filename
