#!/usr/bin/env python3
"""
Generate PILOT_DEMO_OPINION.docx on the host (for P13 smoke test).
This script runs on the host machine, not in the container.
Requires: pip install python-docx
"""
import sys
from pathlib import Path

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    print("[ERROR] python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)

def create_demo_opinion_docx(output_path: str):
    """Create PILOT_DEMO_OPINION.docx with safe placeholder content."""
    doc = DocxDocument()
    
    # Title
    title_para = doc.add_heading("PILOT DEMO OPINION", level=1)
    
    # Content paragraphs (ASCII-only, safe for testing)
    content_lines = [
        "Legal Opinion for Property Transfer",
        "",
        "Property Details:",
        "Plot Number: DEMO-001",
        "Block: B",
        "Phase: 6",
        "Scheme: DHA",
        "",
        "Registry Information:",
        "Registry Number: DEMO-789/2024",
        "Registry Date: 2024-01-20",
        "Registry Office: LDA Lahore",
        "Instrument: Sale Deed",
        "",
        "Party Information:",
        "Borrower: Example Borrower Name",
        "Seller: Example Seller Name",
        "",
        "This is a placeholder legal opinion document for pilot testing.",
        "No real property information is included.",
        "This document is safe to commit to the repository.",
    ]
    
    for line in content_lines:
        para = doc.add_paragraph(line)
        para.paragraph_format.space_after = Pt(6)
    
    # Ensure output directory exists
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    
    # Save document
    doc.save(output_path)
    print(f"[OK] Created: {output_path}")

if __name__ == "__main__":
    # Default output path (relative to repo root)
    if len(sys.argv) > 1:
        output_path = sys.argv[1]
    else:
        # Default to docs/pilot_samples_real_example/PILOT_DEMO_OPINION.docx
        repo_root = Path(__file__).parent.parent.parent
        output_path = str(repo_root / "docs" / "pilot_samples_real_example" / "PILOT_DEMO_OPINION.docx")
    
    create_demo_opinion_docx(output_path)
    print("[OK] DOCX generation complete")

